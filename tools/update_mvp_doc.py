from pathlib import Path
from shutil import copy2
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

root = Path(r"C:\proyectos\mototaxi-atacames")
source = root / "outputs" / "Documento_funcional_MVP_Mototaxi_Atacames.docx"
backup = root / "outputs" / "Documento_funcional_MVP_Mototaxi_Atacames_v1.0_original.docx"
if not backup.exists():
    copy2(source, backup)

doc = Document(source)
for table in doc.tables:
    for row in table.rows:
        if len(row.cells) >= 2 and row.cells[0].text.strip() == "Versión":
            row.cells[1].text = "1.1 · 29 de julio de 2026"

# Avoid duplicating the implementation appendix on repeated runs.
for paragraph in list(doc.paragraphs):
    if paragraph.text.strip() == "13. Estado técnico de la implementación":
        raise SystemExit("La sección 13 ya existe; no se hicieron cambios.")

doc.add_page_break()
doc.add_heading("13. Estado técnico de la implementación", level=1)
p = doc.add_paragraph()
r = p.add_run("Actualización 1.1. ")
r.bold = True
p.add_run("La base técnica y la primera consola administrativa navegable ya están implementadas. Esta sección registra lo construido, cómo reproducir el entorno y qué controles deben reforzarse antes de producción.")

doc.add_heading("13.1 Consola administrativa implementada", level=2)
items = [
    "Inicio de sesión para los roles Administrador y Soporte, con autorización diferenciada en la API.",
    "Tablero operativo con métricas y listado de viajes activos del piloto.",
    "Gestión de conductores: consulta, aprobación, rechazo y suspensión con motivo auditado.",
    "Gestión de pasajeros: consulta, suspensión y reactivación.",
    "Publicación e historial de versiones tarifarias con fecha de vigencia.",
    "Editor poligonal para dibujar zonas urbanas y extendidas y conservar su versión.",
    "Gestión de incidentes: estado, asignación y resolución por soporte.",
    "Auditoría de accesos y cambios sensibles, visible solo para administradores.",
    "Diagnóstico de conexión a PostgreSQL y verificación de la extensión PostGIS."
]
for text in items:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("13.2 Persistencia y entorno reproducible", level=2)
doc.add_paragraph("El repositorio incorpora PostgreSQL 16 con PostGIS 3.4 mediante Docker Compose, migraciones SQL incrementales y un ejecutor que registra cada migración aplicada. La primera migración cubre usuarios, conductores, vehículos, zonas, tarifas, viajes, ofertas, eventos y auditoría. La segunda añade documentos, calificaciones, incidentes y sesiones administrativas.")

table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
hdr = table.rows[0].cells
hdr[0].text, hdr[1].text, hdr[2].text = "Componente", "Comando", "Resultado"
rows = [
    ("Dependencias", "pnpm install", "Instala API, panel y dominio."),
    ("PostGIS", "pnpm db:up", "Inicia PostgreSQL/PostGIS local."),
    ("Migraciones", "pnpm db:migrate", "Crea y versiona el esquema."),
    ("API", "pnpm dev:api", "Servidor en http://localhost:3001."),
    ("Panel", "pnpm dev:admin", "Consola en http://localhost:3000."),
    ("Calidad", "pnpm test", "Ejecuta pruebas del monorepo."),
]
for values in rows:
    cells = table.add_row().cells
    for cell, value in zip(cells, values):
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
for row in table.rows:
    widths = [Inches(1.35), Inches(1.55), Inches(3.6)]
    for cell, width in zip(row.cells, widths):
        cell.width = width
        tcPr = cell._tc.get_or_add_tcPr()
        tcW = tcPr.find(qn("w:tcW"))
        if tcW is None:
            tcW = OxmlElement("w:tcW")
            tcPr.append(tcW)
        tcW.set(qn("w:w"), str(int(width.inches * 1440)))
        tcW.set(qn("w:type"), "dxa")

for cell in table.rows[0].cells:
    for run in cell.paragraphs[0].runs:
        run.bold = True

doc.add_heading("13.3 Credenciales y datos piloto", level=2)
doc.add_paragraph("Para desarrollo local se proporcionan cuentas de demostración configurables mediante variables de entorno:")
for text in [
    "Administrador: admin@mototaxi.local / Mototaxi2026!",
    "Soporte: soporte@mototaxi.local / Soporte2026!"
]:
    doc.add_paragraph(text, style="List Bullet")
warning = doc.add_paragraph()
r = warning.add_run("Advertencia de seguridad. ")
r.bold = True
warning.add_run("Estas credenciales y el secreto de sesión deben sustituirse antes de cualquier piloto con usuarios reales. Los datos visibles actuales son datos piloto en memoria; las tablas PostGIS están listas, pero el repositorio persistente administrativo debe completarse antes de producción.")

doc.add_heading("13.4 Validación realizada", level=2)
for text in [
    "Flutter analyze y pruebas de widget aprobadas.",
    "APK Android de depuración compilado, instalado y ejecutado en un emulador Android 11.",
    "Pruebas de dominio, cotización, autenticación, permisos y auditoría aprobadas.",
    "TypeScript y compilación de producción del panel aprobados.",
    "Entorno Android validado en Windows con Flutter, SDK, JDK, ADB e hipervisor."
]:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("13.5 Pendientes antes de producción", level=2)
for text in [
    "Sustituir el almacén piloto por repositorios PostgreSQL en todas las rutas administrativas.",
    "Persistir usuarios administrativos con contraseñas hasheadas, expiración y revocación de sesiones y MFA.",
    "Integrar proveedor cartográfico real y convertir los puntos del editor a coordenadas geográficas SRID 4326.",
    "Conectar datos reales de viajes, ubicaciones, documentos y evidencias.",
    "Restringir CORS, habilitar HTTPS, límites de intentos, almacenamiento privado y observabilidad.",
    "Completar las validaciones legales, operativas, de privacidad y emergencia indicadas en este documento."
]:
    doc.add_paragraph(text, style="List Bullet")

doc.save(source)
print(source)